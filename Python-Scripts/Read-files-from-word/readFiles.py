from docx.api import Document
import os
import zipfile
import re
import xml.dom.minidom
import csv
import os.path
import sys
from tqdm import tqdm

try:
    from xml.etree.cElementTree import XML
except ImportError:
    from xml.etree.ElementTree import XML
import zipfile

#csvData = [['Title (from file)', 'Location (from file)', 'Date of examination (from file)', 'Component system name (from file)', 'Inspectors name (from file)', 'File name', 'File location on shared folder']]
WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
PARA = WORD_NAMESPACE + 'p'
TEXT = WORD_NAMESPACE + 't'

#filename = "{}.docx".format(input())
#filename = "PL-520 CUI.docx"

filename = []
extensions = {}
source_dir = os.path.dirname(os.path.realpath(__file__))

for dirpath, dirnames, filenames in os.walk(source_dir):
    for dirname in [f for f in dirnames]:
        dirx = os.path.join(dirpath, dirname)
        print(dirx)
        files = os.listdir(dirx)
        for f in files:
            if not f.startswith('~'): # skip file created when file is open
                file_fullpath = dirx + "\\" + f

                if os.path.isdir(file_fullpath):
                    continue

                if f.lower().endswith('.docx') or f.lower().endswith('.docx'):
                    filename.append(file_fullpath)
                    fname, fileExtension = os.path.splitext(f)
                    fileExtension = fileExtension.lower()
                    if fileExtension not in extensions:
                        extensions[fileExtension] = 0
                    extensions[fileExtension] += 1

'''def path_to_file(filename):
    return os.path.abspath(filename)'''

# Files like '2019'
def get_docx_text(path):
    doc = Document(path)

    t = doc.tables[0].cell(0,0).text.strip() #Title name
    l = doc.tables[0].cell(3,6).text.strip() #Location
    d = doc.tables[0].cell(4,12).text.strip() #Date
    s = doc.tables[0].cell(2,3).text.strip() #System name
    i = doc.tables[1].cell(21,0).text.strip() #Inspector
    #p = path[i] #File location

    csvData = [[t, l, d, s, i, filename, path]]
    create_csv(csvData)

# Files like '2018' "H10-PL-718-DA-Rev0"
def get_docx_text2(path):
    doc = Document(path)

    t = "PIPEWORK EXTERNAL VISUAL EXAMINATION REPORT" #Title name
    l = "Hannibal Plant" #Location
    d = doc.tables[0].cell(4,12).text.strip() #Date
    s = doc.tables[0].cell(2,3).text.strip() #System name
    i = doc.tables[4].cell(4,0).text.strip() #Inspector #File location

    csvData = [[t, l, d, s, i, path]]
    create_csv(csvData)

def get_docx_text_try_except(path):
    doc = Document(path)

    t = "PIPEWORK EXTERNAL VISUAL EXAMINATION REPORT" #Title name
    l = "Hannibal Plant" #Location
    try:
        d = doc.tables[0].cell(4,12).text.strip() #Date
    except:
        print("Something wont work date")
    try:
        s = doc.tables[0].cell(2,3).text.strip() #System name
    except:
        print("Something wont work system name")
    try:
        i = doc.tables[4].cell(4,0).text.strip() #Inspector #File location
    except:
        print("Something wont work inspectors name")
    try:
        csvData = [[t, l, d, s, i, path]]
        create_csv(csvData)
    except:
        print("Something wont work")

# Files like '2013'
def get_docx_text3(path):
    doc = Document(path)

    t = "DETAILED VISUAL INSPECTION REPORT" #Title name
    l = doc.tables[0].cell(2,6).text.strip() #Location
    d = doc.tables[0].cell(3,12).text.strip() #Date
    s = doc.tables[0].cell(1,3).text.strip() #System name
    i = doc.tables[3].cell(8,0).text.strip() #Inspector
    #p = path[i] #File location

    csvData = [[t, l, d, s, i, path]]
    create_csv(csvData)

# Files like '2013' "DYE PENETRANT INSPECTION REPORT"
def get_docx_text4(path):
    doc = Document(path)

    t = doc.tables[0].cell(0,1).text.strip() #Title name
    l = doc.tables[0].cell(4,0).text.strip() #Location
    d = doc.tables[0].cell(2,4).text.strip() #Date
    s = doc.tables[0].cell(2,0).text.strip() #System name
    i = doc.tables[1].cell(2,0).text.strip() #Inspector
    #p = path[i] #File location

    csvData = [[t, l, d, s, i, path]]
    create_csv(csvData)

# FIles like '2012'
def get_docx_text5(path):
    doc = Document(path)

    t = "DETAILED VISUAL INSPECTION REPORT" #Title name
    l = doc.tables[0].cell(3,1).text.strip() #Location
    d = doc.tables[0].cell(3,12).text.strip() #Date
    s = doc.tables[0].cell(1,3).text.strip() #System name
    i = doc.tables[3].cell(8,0).text.strip() #Inspector
    #p = path[i] #File location

    csvData = [[t, l, d, s, i, path]]
    create_csv(csvData)

# Files like '2007'
def get_docx_text6(path):
    doc = Document(path)

    t = doc.tables[0].cell(0,0).text.strip() #Title name
    l = "Hannibal Plant" #Location
    d = doc.tables[1].cell(5,0).text.strip() #Date
    s = doc.tables[1].cell(1,5).text.strip() #System name
    i = doc.tables[7].cell(1,0).text.strip() #Inspector
    #p = path[i] #File location

    csvData = [[t, l, d, s, i, path]]
    create_csv(csvData)

# Files like '2006'
def get_docx_text7(path):
    doc = Document(path)

    t = doc.tables[0].cell(0,0).text.strip() #Title name
    l = "Hannibal Plant" #Location
    d = doc.tables[1].cell(5,0).text.strip() #Date
    s = doc.tables[1].cell(1,3).text.strip() #System name
    i = doc.tables[4].cell(7,0).text.strip() #Inspector
    #p = path[i] #File location

    csvData = [[t, l, d, s, i, path]]
    create_csv(csvData)

def create_csv(csvData):
    with open('Piping-inspections.csv', 'a') as csvFile:
        writer = csv.writer(csvFile)
        writer.writerows(csvData)

    csvFile.close()

for i in tqdm(range(0, len(filename))):
    try:
        print(filename[i])
        get_docx_text7(filename[i])

    except Exception as e:
        print(e)

os.system("PAUSE")
